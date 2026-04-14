#!/usr/bin/env python3
"""
md-to-docx.py — Convert a resume Markdown file to a professional, ATS-friendly .docx file.

Usage:
    python md-to-docx.py <path/to/Resume.md> [output/dir] [author_name]

    If output/dir is omitted, saves alongside the input file.
    If output/dir is provided, saves there with the same stem.
    If author_name is omitted, name is auto-detected from the H1 (# Name) line.
    The author_name is written to .docx core properties (Author, Last Modified By).

Requirements:
    pip install python-docx"""

import re
import sys
from pathlib import Path


def check_dependency():
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("ERROR: python-docx is not installed.\nRun:  pip install python-docx")
        sys.exit(1)


def strip_markdown_inline(text: str) -> str:
    """Strip inline markdown, keep plain text (used for non-hyperlink contexts)."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text.strip()


def extract_links(text: str):
    """Return list of (display_text, url) tuples found in markdown link syntax."""
    return re.findall(r'\[([^\]]+)\]\(([^)]+)\)', text)


def extract_bold_pairs(text: str):
    """
    For a skills line like '**Category:** skill1, skill2'
    returns ('Category:', 'skill1, skill2').
    Otherwise returns (None, plain_text).
    """
    m = re.match(r'\*\*(.+?):\*\*\s*(.*)', text)
    if m:
        return m.group(1) + ':', m.group(2).strip()
    return None, strip_markdown_inline(text)


def add_hyperlink(paragraph, display_text: str, url: str, font_size, color_rgb):
    """Insert a proper clickable hyperlink run into a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    # Add relationship
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri Light')
    rPr.append(rFonts)

    # Font size (half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)

    # Colour
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color_rgb))
    rPr.append(color)

    # Underline style (u element)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def detect_author_from_md(md_path: str) -> str:
    """Extract the H1 name line as author fallback."""
    with open(md_path, encoding='utf-8') as f:
        for line in f:
            stripped = line.strip()
            if stripped.startswith('# '):
                return stripped[2:].strip()
    return ''


def build_docx(md_path: str, output_dir: str = None, author_name: str = None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins (ATS safe: generous, clean) ───────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Default paragraph spacing ─────────────────────────────────────────────
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(2)
    style.font.name = 'Calibri Light'
    style.font.size = Pt(11)

    DARK      = RGBColor(0x00, 0x00, 0x00)
    ACCENT    = RGBColor(0x1F, 0x4E, 0x79)   # dark navy for headings/name
    GREY      = RGBColor(0x55, 0x55, 0x55)
    LINK_CLR  = RGBColor(0x1F, 0x4E, 0x79)   # same navy for hyperlinks

    def add_name(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = ACCENT

    def add_contact(text):
        """
        Render the contact line, converting any [display](url) to real hyperlinks.
        Plain text segments separated by | are rendered as-is.
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)

        # Split line into segments around markdown links
        link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        segments = link_pattern.split(text)
        # split produces: [before, display, url, between, display, url, after, ...]

        idx = 0
        while idx < len(segments):
            chunk = segments[idx]
            if idx % 3 == 0:
                # Plain text segment — strip ** and *
                plain = re.sub(r'\*\*(.+?)\*\*', r'\1', chunk)
                plain = re.sub(r'\*(.+?)\*', r'\1', plain)
                if plain:
                    run = p.add_run(plain)
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = DARK
            elif idx % 3 == 1:
                # display text for hyperlink
                display = segments[idx]
                url     = segments[idx + 1]
                add_hyperlink(p, display, url, 9.5, (0x1F, 0x4E, 0x79))
                idx += 1  # skip the url segment
            idx += 1

    def add_section_heading(text):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT
        # Bottom border underline rule
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E79')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_role_header(title_company):
        """Bold role title line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(strip_markdown_inline(title_company))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

    def add_date_location(text):
        """Italic date | location line beneath role header."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(strip_markdown_inline(text))
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent   = Inches(0.2)
        p.paragraph_format.space_before  = Pt(1.5)
        p.paragraph_format.space_after   = Pt(1.5)
        run = p.add_run(strip_markdown_inline(text))
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

    def add_skill_line(text):
        """Bold category label followed by plain skill list on same line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        label, rest = extract_bold_pairs(text)
        if label:
            r1 = p.add_run(label + ' ')
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = DARK
            r2 = p.add_run(rest)
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK
        else:
            run = p.add_run(rest)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    def add_tech_stack(text):
        """Small italic tech stack line after role bullets."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(strip_markdown_inline(text))
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY

    def add_right_to_work(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(0)
        label, rest = extract_bold_pairs(text)
        if label:
            r1 = p.add_run(label + ' ')
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(rest)
            r2.font.size = Pt(10)
        else:
            p.add_run(strip_markdown_inline(text)).font.size = Pt(10)

    def add_body_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(strip_markdown_inline(text))
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

    # ── Parse and render ──────────────────────────────────────────────────────
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()

    i = 0
    name_done = False
    contact_done = False
    current_h3 = None     # accumulate H3 title across lines
    pending_date = None   # date line after an H3

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Skip horizontal rules
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # H1 → candidate name
        if stripped.startswith('# ') and not name_done:
            add_name(stripped[2:].strip())
            name_done = True
            i += 1
            continue

        # Contact line (bold + pipe pattern, right after name)
        if name_done and not contact_done and stripped.startswith('**') and '|' in stripped:
            add_contact(stripped)
            contact_done = True
            i += 1
            continue

        # H2 → section heading
        if stripped.startswith('## '):
            add_section_heading(stripped[3:].strip())
            i += 1
            current_h3 = None
            continue

        # H3 → role / project / edu heading
        if stripped.startswith('### '):
            current_h3 = stripped[4:].strip()
            i += 1
            # Peek at next non-empty line to see if it's a date/location
            j = i
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                next_line = lines[j].strip()
                # Date/location lines start with ** and contain YYYY or –
                if re.match(r'\*\*.+\*\*', next_line) and ('20' in next_line or '–' in next_line or 'Present' in next_line):
                    add_role_header(current_h3)
                    add_date_location(next_line)
                    i = j + 1
                    current_h3 = None
                    continue
            add_role_header(current_h3)
            current_h3 = None
            continue

        # Bold date/location line (after H3 consumed above, or standalone)
        if re.match(r'^\*\*.+\*\*.*\|', stripped) and '20' in stripped:
            add_date_location(stripped)
            i += 1
            continue

        # Bullet point
        if stripped.startswith('- '):
            add_bullet(stripped[2:].strip())
            i += 1
            continue

        # Tech stack italic line (*Tech Stack: ...*)
        if stripped.startswith('*') and stripped.endswith('*') and 'Tech Stack' in stripped:
            add_tech_stack(stripped.strip('*').strip())
            i += 1
            continue

        # Right to work line
        if stripped.startswith('**Right to Work') or 'Stamp 4' in stripped:
            add_right_to_work(stripped)
            i += 1
            continue

        # Skill line with bold category
        if stripped.startswith('**') and ':**' in stripped:
            add_skill_line(stripped)
            i += 1
            continue

        # Non-empty text → body paragraph
        if stripped:
            add_body_text(stripped)

        i += 1

    # ── Save ──────────────────────────────────────────────────────────────────
    input_path = Path(md_path)
    if output_dir:
        out_path = Path(output_dir) / (input_path.stem + '.docx')
    else:
        out_path = input_path.with_suffix('.docx')

    out_path.parent.mkdir(parents=True, exist_ok=True)

    # ── Set .docx author metadata ─────────────────────────────────────────────
    resolved_author = author_name or detect_author_from_md(md_path)
    if resolved_author:
        doc.core_properties.author = resolved_author
        doc.core_properties.last_modified_by = resolved_author

    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    if resolved_author:
        print(f"Author: {resolved_author}")
    return str(out_path)


def build_docx_from_string(md_content: str, output_path: str, author_name: str = None):
    """Build .docx from a markdown string (no intermediate .md file persisted).

    Used by tailor-resume.py to generate .docx directly from in-memory content.
    """
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
        tmp.write(md_content)
        tmp_path = tmp.name
    try:
        out_dir = str(Path(output_path).parent)
        Path(out_dir).mkdir(parents=True, exist_ok=True)
        build_docx(tmp_path, out_dir, author_name)
        temp_docx = Path(out_dir) / (Path(tmp_path).stem + '.docx')
        final_out = Path(output_path)
        if temp_docx.exists() and temp_docx.resolve() != final_out.resolve():
            if final_out.exists():
                final_out.unlink()
            temp_docx.rename(final_out)
    finally:
        Path(tmp_path).unlink(missing_ok=True)
    print(f"Saved: {output_path}")
    if author_name:
        print(f"Author: {author_name}")
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python md-to-docx.py <Resume.md> [output/dir] [author_name]")
        sys.exit(1)

    md_path    = sys.argv[1]
    output_dir  = sys.argv[2] if len(sys.argv) > 2 else None
    author_name = sys.argv[3] if len(sys.argv) > 3 else None

    if not Path(md_path).exists():
        print(f"ERROR: File not found: {md_path}")
        sys.exit(1)

    check_dependency()
    build_docx(md_path, output_dir, author_name)


if __name__ == "__main__":
    main()
